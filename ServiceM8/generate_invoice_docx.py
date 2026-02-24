
import sys
import os
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: python-docx is not installed. Please install it with: pip install python-docx")
    sys.exit(1)

def set_cell_background(cell, color_hex):
    """
    Helper to set cell background color.
    color_hex: string like "0F172A" (no hash)
    """
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_merge_field(paragraph, field_name, bold=False, color=None, size=None):
    """
    Helper to insert a MERGEFIELD.
    """
    run = paragraph.add_run()
    r = run._r
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar_begin)
    
    run = paragraph.add_run()
    r = run._r
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' MERGEFIELD  "{field_name}"  \\* MERGEFORMAT '
    r.append(instrText)
    
    run = paragraph.add_run()
    r = run._r
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    r.append(fldChar_sep)
    
    # Display Text
    display_run = paragraph.add_run(f'«{field_name}»')
    if bold: display_run.bold = True
    if color: display_run.font.color.rgb = color
    if size: display_run.font.size = size
    
    run = paragraph.add_run()
    r = run._r
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r.append(fldChar_end)
    return display_run

def create_invoice(output_path, logo_path):
    doc = Document()
    
    # --- COLORS & FONTS ---
    NAVY = RGBColor(0x0F, 0x17, 0x2A)   # #0F172A
    ORANGE = RGBColor(0xFF, 0x7D, 0x00) # #FF7D00
    GRAY = RGBColor(0x64, 0x74, 0x8B)   # #64748B
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # --- HEADER (Image Based) ---
    # Insert the pre-generated header image at top margin
    section = doc.sections[0]
    section.top_margin = Inches(0.5) 
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # We want the image to span the full width (approx 7.5 inches for A4 with margins)
    # A4 width is 8.27 inches. If margins are 0.5, content width is 7.27.
    # We can overshoot margins if we want "bleed" effect by setting negative indentation, 
    # but standard image insertion is safer.
    
    header_path = "header_bg.png"
    if os.path.exists(header_path):
        # Add to document body at the very top
        # (Alternatively, add to actual header section if it should repeat on pages)
        # For an invoice (usually 1 page), body is fine and easier to control positioning.
        
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(header_path, width=Inches(7.5))
    else:
        print("Header image not found at " + header_path)
        doc.add_heading("TAX INVOICE", 0)

    # Add Job ID below header (Right Aligned) to match the visual flow
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Job Ref: ")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    add_merge_field(p, "job.generated_job_id", bold=True, color=NAVY)

    # Add some spacing below header
    doc.add_paragraph()

    # --- INFO GRID (4 Columns) ---
    info_table = doc.add_table(rows=2, cols=4)
    info_table.width = Inches(7.5)
    
    # Labels
    labels = ["DATE ISSUED", "PO NUMBER", "DUE DATE", "AMOUNT DUE"]
    for i, label in enumerate(labels):
        cell = info_table.cell(0, i)
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY
        run.bold = True
    
    # Values
    # Date
    c = info_table.cell(1, 0)
    add_merge_field(c.paragraphs[0], "job.invoice_date_extended", bold=True, color=NAVY)
    
    # PO
    c = info_table.cell(1, 1)
    add_merge_field(c.paragraphs[0], "job.purchase_order_number", bold=True, color=NAVY)
    
    # Due Date
    c = info_table.cell(1, 2)
    run = c.paragraphs[0].add_run("Upon Receipt")
    run.bold = True
    run.font.color.rgb = NAVY
    
    # Amount Due (Accent Color)
    c = info_table.cell(1, 3)
    add_merge_field(c.paragraphs[0], "job.balance_due", bold=True, color=ORANGE)

    doc.add_paragraph()
    
    # --- ADDRESSES ---
    addr_table = doc.add_table(rows=1, cols=2)
    addr_table.width = Inches(7.5)
    
    # FROM (Static Details from original invoice)
    c1 = addr_table.cell(0, 0)
    p = c1.paragraphs[0]
    run = p.add_run("FROM")
    run.font.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(11)
    p.add_run("\n")
    
    add_merge_field(p, "vendor.name", bold=True, color=NAVY)
    p.add_run("\n52 Manor Farm Road\nSouthampton, Hampshire SO18")
    p.add_run("\n")
    add_merge_field(p, "location.phone_1")
    p.add_run("\n")
    add_merge_field(p, "vendor.email")
    p.add_run("\n")
    add_merge_field(p, "vendor.website")

    # BILL TO
    c2 = addr_table.cell(0, 1)
    p = c2.paragraphs[0]
    run = p.add_run("INVOICE TO")
    run.font.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(11)
    p.add_run("\n")
    # Content
    add_merge_field(p, "job.company_name", bold=True, color=NAVY)
    p.add_run("\n")
    add_merge_field(p, "job.billing_address")
    
    p.add_run("\n\n")
    run = p.add_run("JOB ADDRESS")
    run.font.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(11)
    p.add_run("\n")
    add_merge_field(p, "job.contact_first")
    p.add_run(" ")
    add_merge_field(p, "job.contact_last")
    p.add_run("\n")
    add_merge_field(p, "job.job_address")
    
    doc.add_paragraph()

    # --- WORK DESCRIPTION ---
    doc.add_heading('Work Carried Out', level=3).runs[0].font.color.rgb = NAVY
    p = doc.add_paragraph()
    add_merge_field(p, "job.work_done_description")
    
    doc.add_paragraph()

    # --- ITEMS TABLE ---
    items_table = doc.add_table(rows=2, cols=4)
    items_table.style = 'Table Grid'
    items_table.width = Inches(7.5)
    
    # Header Row
    headers = ["Description", "Price", "Quantity", "Total"]
    row0 = items_table.rows[0]
    # Set blue header background
    for i, h in enumerate(headers):
        cell = row0.cells[i]
        set_cell_background(cell, "0F172A")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10) # Slightly larger
        if i > 1: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if i == 1: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Price right aligned

    # Data Row
    row1 = items_table.rows[1]
    
    # Col 0: Desc
    add_merge_field(row1.cells[0].paragraphs[0], "jobMaterial.name", bold=True, color=NAVY)
    
    # Col 1: Price
    p = row1.cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_merge_field(p, "jobMaterial.price_ex_tax")
    
    # Col 2: Qty
    p = row1.cells[2].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Center Qty
    add_merge_field(p, "jobMaterial.quantity")
    
    # Col 3: Total
    p = row1.cells[3].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_merge_field(p, "jobMaterial.total_price_ex_tax", bold=True, color=NAVY)

    doc.add_paragraph()

    # --- TOTALS ---
    # We want this right aligned.
    totals_table = doc.add_table(rows=5, cols=2)
    totals_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    
    def add_total_row(idx, label, field, bold_val=False, color_val=None, size_val=None):
        c1 = totals_table.cell(idx, 0)
        c2 = totals_table.cell(idx, 1)
        c1.text = label
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        p = c2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_merge_field(p, field, bold=bold_val, color=color_val, size=size_val)
        
        if bold_val:
            c1.paragraphs[0].runs[0].bold = True
        if color_val:
             c1.paragraphs[0].runs[0].font.color.rgb = color_val
        if size_val:
             c1.paragraphs[0].runs[0].font.size = size_val
             
    add_total_row(0, "Sub Total:", "job.subtotal_price")
    add_total_row(1, "VAT:", "job.total_tax_price")
    add_total_row(2, "TOTAL:", "job.total_price", bold_val=True, color_val=NAVY, size_val=Pt(14))
    add_total_row(3, "Payments:", "job.amount_paid", color_val=RGBColor(0x10, 0xB9, 0x81)) # Green
    add_total_row(4, "BALANCE DUE:", "job.balance_due", bold_val=True, color_val=RGBColor(0xEF, 0x44, 0x44)) # Red

    doc.add_paragraph()
    
    # --- FOOTER (Bank Details) ---
    section = doc.sections[0]
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add text
    run = footer_p.add_run("Better Call Wes Plumbing & Heating")
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(12)
    
    footer_p.add_run("\n\nPAYMENT DETAILS")
    
    run = footer_p.add_run("\nBank Transfer:")
    run.bold = True
    
    footer_p.add_run("\nAccount Name: Better Call Wes Plumbing & Heating")
    footer_p.add_run("\nBank Name: Monzo | Sort Code: 04-00-03 | Account No: 58272294")
    
    footer_p.add_run("\n\nWe Accept: Bank Transfer, Card (Visa, Mastercard)")
    
    doc.save(output_path)
    print(f"Successfully generated {output_path}")

if __name__ == "__main__":
    LOGO_PATH = "/Users/akweteybortier/Coding/Better Call Wes/Brand Images/a1781cea-8d5b-423d-bd16-14e4f373f63c.png"
    OUTPUT_PATH = "BCW - Invoice Final.docx"
    
    create_invoice(OUTPUT_PATH, LOGO_PATH)
